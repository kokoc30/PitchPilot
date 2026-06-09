from io import BytesIO
from pathlib import Path

from app.services.resume_chunker import clean_resume_text
from app.utils.config import get_settings

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


class ResumeParseError(Exception):
    def __init__(self, message: str, *, code: str = "resume_parse_error") -> None:
        super().__init__(message)
        self.code = code
        self.message = message


def parse_resume_file(
    filename: str,
    content_type: str | None,
    content: bytes,
) -> str:
    settings = get_settings()
    max_bytes = settings.resume_max_file_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise ResumeParseError(
            f"Resume file is larger than {settings.resume_max_file_mb} MB.",
            code="resume_file_too_large",
        )

    file_type = detect_resume_file_type(filename, content_type)
    if file_type == "txt":
        text = _parse_txt(content)
    elif file_type == "pdf":
        text = _parse_pdf(content)
    elif file_type == "docx":
        text = _parse_docx(content)
    else:
        raise ResumeParseError(
            "Unsupported resume type. Upload a PDF, DOCX, or TXT file.",
            code="unsupported_resume_type",
        )

    cleaned = clean_resume_text(text)
    if not cleaned:
        raise ResumeParseError(
            "No readable resume text was extracted from this file.",
            code="empty_resume_text",
        )
    return cleaned


def detect_resume_file_type(filename: str, content_type: str | None) -> str:
    suffix = Path(filename or "").suffix.lower()
    normalized_content_type = (content_type or "").split(";")[0].strip().lower()

    if suffix == ".pdf" or normalized_content_type == "application/pdf":
        return "pdf"
    if (
        suffix == ".docx"
        or normalized_content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return "docx"
    if suffix == ".txt" or normalized_content_type == "text/plain":
        return "txt"

    raise ResumeParseError(
        "Unsupported resume type. Upload a PDF, DOCX, or TXT file.",
        code="unsupported_resume_type",
    )


def _parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeParseError(
            "PDF parsing dependency is not installed.",
            code="resume_dependency_missing",
        ) from exc

    try:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ResumeParseError(
            "Unable to extract text from this PDF.",
            code="pdf_extraction_failed",
        ) from exc

    return "\n\n".join(pages)


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeParseError(
            "DOCX parsing dependency is not installed.",
            code="resume_dependency_missing",
        ) from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ResumeParseError(
            "Unable to read this DOCX resume.",
            code="docx_extraction_failed",
        ) from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)
